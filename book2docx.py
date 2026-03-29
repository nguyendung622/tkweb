#!/usr/bin/env python3
"""
Convert "Làm báo trên môi trường số" HTML book to DOCX
"""

import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BOOK_DIR = os.path.join(os.path.dirname(__file__), "book")

CHAPTERS = [
    ("chapters/00-loi-noi-dau.html",        "Lời nói đầu"),
    ("chapters/00-huong-dan-su-dung.html",   "Hướng dẫn sử dụng sách"),
    ("chapters/01-google-sites.html",        "Chương 1"),
    ("chapters/02-ux-co-ban.html",           "Chương 2"),
    ("chapters/03-canva.html",               "Chương 3"),
    ("chapters/04-wordpress.html",           "Chương 4"),
    ("chapters/05-seo.html",                 "Chương 5"),
    ("chapters/06-du-lieu.html",             "Chương 6"),
    ("chapters/07-mang-xa-hoi.html",         "Chương 7"),
    ("chapters/08-multimedia.html",          "Chương 8"),
    ("chapters/09-bao-mat.html",             "Chương 9"),
    ("chapters/10-newsletter.html",          "Chương 10"),
    ("chapters/11-analytics.html",           "Chương 11"),
    ("chapters/12-du-an.html",               "Chương 12"),
    ("appendix/glossary.html",               "Phụ lục: Bảng thuật ngữ"),
    ("appendix/references.html",             "Phụ lục: Tài liệu tham khảo"),
]

# ─── Màu sắc ─────────────────────────────────────────────
COLOR_ACCENT   = RGBColor(0x1a, 0x56, 0xdb)   # xanh chính
COLOR_NOTE_BG  = RGBColor(0xef, 0xf6, 0xff)
COLOR_WARN_BG  = RGBColor(0xff, 0xf7, 0xed)
COLOR_TIP_BG   = RGBColor(0xf0, 0xfd, 0xf4)
COLOR_MUTED    = RGBColor(0x6b, 0x72, 0x80)
COLOR_HEADING2 = RGBColor(0x11, 0x18, 0x27)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_border(table, size=4, color="D1D5DB"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_run_text(para, text, bold=False, italic=False, color=None, size=None, code=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return run


def para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def process_inline(run_adder, element):
    """Add inline content (text, bold, italic, code, links) to a paragraph via run_adder callback."""
    for child in element.children:
        if isinstance(child, str):
            text = child
            if text:
                run_adder(text)
        elif child.name in ("strong", "b"):
            run_adder(child.get_text(), bold=True)
        elif child.name in ("em", "i"):
            run_adder(child.get_text(), italic=True)
        elif child.name == "code":
            run_adder(child.get_text(), code=True)
        elif child.name == "a":
            run_adder(child.get_text())
        elif child.name in ("span", "mark"):
            process_inline(run_adder, child)
        else:
            text = child.get_text()
            if text.strip():
                run_adder(text)


def add_paragraph_from_element(doc, element, style="Normal"):
    para = doc.add_paragraph(style=style)
    para_spacing(para, after=8)
    def adder(text, bold=False, italic=False, code=False):
        add_run_text(para, text, bold=bold, italic=italic, code=code)
    process_inline(adder, element)
    return para


def add_box(doc, box_el, box_type):
    """Render callout boxes as a shaded single-cell table."""
    icons = {"note": "📌", "tip": "💡", "warning": "⚠️", "practice": "🛠"}
    bg_colors = {
        "note":     RGBColor(0xef, 0xf6, 0xff),
        "tip":      RGBColor(0xf0, 0xfd, 0xf4),
        "warning":  RGBColor(0xff, 0xf7, 0xed),
        "practice": RGBColor(0xf5, 0xf3, 0xff),
    }
    bg = bg_colors.get(box_type, RGBColor(0xf9, 0xfa, 0xfb))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, size=4, color="E5E7EB")
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)

    # Title
    title_el = box_el.find(class_="box-title")
    title_text = title_el.get_text(strip=True) if title_el else icons.get(box_type, "")
    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after = Pt(4)
    add_run_text(tp, title_text, bold=True, size=9.5)

    # Content
    for child in box_el.children:
        if hasattr(child, "get") and child.get("class") and "box-title" in child.get("class", []):
            continue
        if child.name == "p":
            cp = cell.add_paragraph()
            cp.paragraph_format.space_after = Pt(4)
            def adder(text, bold=False, italic=False, code=False):
                add_run_text(cp, text, bold=bold, italic=italic, code=code, size=9.5)
            process_inline(adder, child)
        elif child.name in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                lp = cell.add_paragraph(style="List Bullet")
                lp.paragraph_format.space_after = Pt(2)
                def adder(text, bold=False, italic=False, code=False):
                    add_run_text(lp, text, bold=bold, italic=italic, code=code, size=9.5)
                process_inline(adder, li)

    # Spacing after box
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, table_el):
    rows_el = table_el.find_all("tr")
    if not rows_el:
        return
    max_cols = max(len(r.find_all(["th", "td"])) for r in rows_el)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows_el), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, size=4, color="D1D5DB")

    for r_idx, row_el in enumerate(rows_el):
        cells_el = row_el.find_all(["th", "td"])
        is_header = row_el.find("th") is not None
        for c_idx, cell_el in enumerate(cells_el):
            if c_idx >= max_cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            if is_header:
                set_cell_bg(cell, RGBColor(0xf3, 0xf4, 0xf6))
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            def adder(text, bold=False, italic=False, code=False, _is_header=is_header):
                add_run_text(cp, text, bold=(_is_header or bold), italic=italic, code=code, size=9.5)
            process_inline(adder, cell_el)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_list(doc, list_el, ordered=False, level=0):
    style = "List Number" if ordered else "List Bullet"
    for li in list_el.find_all("li", recursive=False):
        # Check for nested lists
        nested = li.find(["ul", "ol"])
        # Text part (exclude nested list text)
        text_parts = []
        for child in li.children:
            if child.name in ("ul", "ol"):
                continue
            text_parts.append(child)

        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(3)
        if level > 0:
            para.paragraph_format.left_indent = Cm(level * 0.5)

        def adder(text, bold=False, italic=False, code=False):
            add_run_text(para, text, bold=bold, italic=italic, code=code)

        for part in text_parts:
            if isinstance(part, str):
                if part.strip():
                    adder(part)
            elif hasattr(part, "name"):
                process_inline(adder, part)

        if nested:
            add_list(doc, nested, ordered=(nested.name == "ol"), level=level+1)


def process_chapter(doc, html_path):
    with open(html_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    main = soup.find("main")
    if not main:
        return

    # Chapter header
    header = main.find(class_="chapter-header")
    if header:
        num_el = header.find(class_="chapter-num")
        h1_el = header.find("h1")
        if num_el:
            p = doc.add_paragraph(num_el.get_text(strip=True))
            p.style = "Heading 3"
            p.runs[0].font.color.rgb = COLOR_MUTED
        if h1_el:
            p = doc.add_heading(h1_el.get_text(strip=True), level=1)
            p.runs[0].font.color.rgb = COLOR_ACCENT

    # Learning objectives
    lo = main.find(class_="learning-objectives")
    if lo:
        p = doc.add_paragraph()
        add_run_text(p, "Mục tiêu học tập", bold=True, size=10)
        para_spacing(p, after=3)
        for li in lo.find_all("li"):
            lp = doc.add_paragraph(style="List Bullet")
            lp.paragraph_format.space_after = Pt(2)
            add_run_text(lp, li.get_text(strip=True), size=9.5)

    # Main content — skip nav, header, learning-objectives, chapter-nav-bar, review-questions handled separately
    skip_classes = {"chapter-header", "learning-objectives", "chapter-nav-bar", "book-nav", "site-title", "book-title"}

    content_area = main
    for element in content_area.children:
        if not hasattr(element, "name") or not element.name:
            continue
        el_classes = set(element.get("class", []))
        if el_classes & skip_classes:
            continue
        if element.name == "nav":
            continue

        if element.name == "h2":
            p = doc.add_heading(element.get_text(strip=True), level=2)
            if p.runs:
                p.runs[0].font.color.rgb = COLOR_HEADING2

        elif element.name == "h3":
            p = doc.add_heading(element.get_text(strip=True), level=3)

        elif element.name == "p":
            add_paragraph_from_element(doc, element)

        elif element.name == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(element.get_text(" ", strip=True))
            run.italic = True
            run.font.color.rgb = COLOR_MUTED

        elif element.name == "ul":
            add_list(doc, element, ordered=False)

        elif element.name == "ol":
            add_list(doc, element, ordered=True)

        elif element.name == "table":
            add_table(doc, element)

        elif element.name == "figure":
            cap = element.find("figcaption")
            if cap:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_run_text(p, cap.get_text(strip=True), italic=True, color=COLOR_MUTED, size=9)

        elif element.name == "div":
            # Callout boxes
            if "box" in el_classes:
                box_type = "note"
                for bt in ("note", "tip", "warning", "practice"):
                    if bt in el_classes:
                        box_type = bt
                        break
                add_box(doc, element, box_type)

            # review-questions
            elif "review-questions" in el_classes:
                p = doc.add_heading("Câu hỏi ôn tập", level=2)
                ol = element.find("ol")
                if ol:
                    add_list(doc, ol, ordered=True)

            # TOC grid (index page)
            elif "toc-grid" in el_classes:
                for item in element.find_all(class_="toc-item"):
                    num = item.find(class_="toc-num")
                    title = item.find(class_="toc-title")
                    dur = item.find(class_="toc-duration")
                    if title:
                        p = doc.add_paragraph(style="List Bullet")
                        num_text = (num.get_text(strip=True) + ". ") if num and num.get_text(strip=True) != "—" else ""
                        add_run_text(p, num_text + title.get_text(strip=True), bold=bool(num_text))
                        if dur and dur.get_text(strip=True):
                            add_run_text(p, f"  ({dur.get_text(strip=True)})", color=COLOR_MUTED, size=9)

            # Author block (index cover)
            elif element.find(class_=None) is None or True:
                text = element.get_text(strip=True)
                if text:
                    # recurse into generic divs for content
                    for child in element.children:
                        if not hasattr(child, "name") or not child.name:
                            continue
                        child_classes = set(child.get("class", []))
                        if child.name == "h2":
                            doc.add_heading(child.get_text(strip=True), level=2)
                        elif child.name == "h3":
                            doc.add_heading(child.get_text(strip=True), level=3)
                        elif child.name == "p":
                            add_paragraph_from_element(doc, child)
                        elif child.name == "ul":
                            add_list(doc, child)
                        elif child.name == "ol":
                            add_list(doc, child, ordered=True)
                        elif child.name == "table":
                            add_table(doc, child)


def setup_styles(doc):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(18)

    # Heading 1 — chapter title
    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_ACCENT
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(10)

    # Heading 2 — section
    h2 = styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_HEADING2
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 — subsection
    h3 = styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # List Bullet
    lb = styles["List Bullet"]
    lb.font.name = "Times New Roman"
    lb.font.size = Pt(11)
    lb.paragraph_format.space_after = Pt(3)

    # List Number
    ln = styles["List Number"]
    ln.font.name = "Times New Roman"
    ln.font.size = Pt(11)
    ln.paragraph_format.space_after = Pt(3)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Trường Đại học Khoa học, Đại học Huế")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_MUTED

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("Làm báo trên môi trường số")
    run2.font.name = "Arial"
    run2.font.size = Pt(32)
    run2.font.bold = True
    run2.font.color.rgb = COLOR_ACCENT

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(40)
    run3 = p3.add_run("Giáo trình dành cho sinh viên ngành Báo chí & Truyền thông")
    run3.font.name = "Arial"
    run3.font.size = Pt(13)
    run3.font.color.rgb = COLOR_MUTED
    run3.font.italic = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(4)
    run4 = p4.add_run("Nguyễn Dũng")
    run4.font.name = "Arial"
    run4.font.size = Pt(14)
    run4.font.bold = True

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_after = Pt(2)
    run5 = p5.add_run("Giảng viên Khoa Công nghệ thông tin")
    run5.font.size = Pt(11)
    run5.font.color.rgb = COLOR_MUTED

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_after = Pt(60)
    run6 = p6.add_run("Trường Đại học Khoa học, Đại học Huế")
    run6.font.size = Pt(11)
    run6.font.color.rgb = COLOR_MUTED

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run7 = p7.add_run("Huế, 2026")
    run7.font.size = Pt(11)
    run7.font.color.rgb = COLOR_MUTED


def main():
    doc = Document()

    # Page setup: A4, margins
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

    setup_styles(doc)
    add_cover(doc)

    for rel_path, label in CHAPTERS:
        html_path = os.path.join(BOOK_DIR, rel_path)
        if not os.path.exists(html_path):
            print(f"  [skip] {rel_path} — không tìm thấy")
            continue

        print(f"  [ok]   {label} — {rel_path}")
        doc.add_page_break()
        process_chapter(doc, html_path)

    out_path = os.path.join(os.path.dirname(__file__), "Lam-bao-tren-moi-truong-so.docx")
    doc.save(out_path)
    print(f"\n✓ Đã lưu: {out_path}")


if __name__ == "__main__":
    main()
